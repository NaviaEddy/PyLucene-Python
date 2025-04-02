import os
import fitz
import pytesseract
from PIL import Image
from docx import Document as DocxDocument
import pandas as pd
import lucene
import psycopg2
from flask import Flask, request, jsonify, render_template, redirect, url_for, flash
from java.io import File
from org.apache.lucene.analysis.standard import StandardAnalyzer
from org.apache.lucene.document import Document, Field, TextField
from org.apache.lucene.index import DirectoryReader, IndexWriter, IndexWriterConfig
from org.apache.lucene.queryparser.classic import QueryParser
from org.apache.lucene.search import IndexSearcher
from org.apache.lucene.store import FSDirectory

# `app = Flask(__name__)` crea una instancia de aplicación Flask. El constructor `Flask(__name__)` inicializa un objeto # Flask que representa la aplicación Flask.
# inicializa un objeto Flask, que representa la aplicación Flask. La variable `__name__` es una
# variable especial de Python que representa el nombre del módulo actual. Cuando se pasa `__name__` al
# al constructor de Flask, ayuda a Flask a determinar la ruta raíz de la aplicación.
app = Flask(__name__)
# La línea `app.secret_key = «tu_clave_secreta»` en el código de la aplicación Flask está estableciendo una clave secreta
# para la aplicación Flask.
app.secret_key = "tu_clave_secreta"  
# La línea `INDEX_DIR = «index_dir»` está definiendo una variable llamada `INDEX_DIR` y asignándole el
# valor `«index_dir»`. Esta variable se utiliza para almacenar la ruta del directorio donde los archivos de índice para el
# función de búsqueda. En este caso, los ficheros índice se almacenarán en un directorio # llamado «directorio_índice».
# llamado «index_dir».
INDEX_DIR = "index_dir"
# La línea `ALLOWED_EXTENSIONS = {«txt», «docx», «pdf», «png», «jpg», «jpeg», «xls», «xlsx»}` es
# definir un conjunto llamado `ALLOWED_EXTENSIONS` que contiene las extensiones de archivo permitidas para el procesamiento
# dentro de la aplicación. Estas extensiones representan los tipos de archivos que pueden ser manejados por la # aplicación para tareas como la extracción de texto.
# para tareas como la extracción de texto, la indexación y la búsqueda.
ALLOWED_EXTENSIONS = {"txt", "docx", "pdf", "png", "jpg", "jpeg",  "xls", "xlsx"}

# Inicializar la máquina virtual de Java (PyLucene)
lucene.initVM()

def allowed_file(filename):
    """
    La función `allowed_file` comprueba si un nombre de archivo dado tiene una extensión de archivo permitida basada en una
    lista predefinida de extensiones permitidas.
    
    :param nombre_archivo: El parámetro `nombre_archivo` es una cadena que representa el nombre de un archivo
    :return: La función `archivo_permitido` devuelve un valor booleano. Comprueba si la entrada `filename`
    contiene un punto (".") y si la extensión del archivo (la parte después del último punto) está en la lista de
    `EXTENSIONES_PERMITIDAS`.
    """
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_txt(file):
    """
    La función `extract_text_from_txt` lee un archivo de texto y devuelve su contenido como una cadena UTF-8 decodificada
    decodificado.
    
    param archivo: Se espera que el parámetro `file` de la función `extract_text_from_txt` sea un archivo
    del que se pueda leer. Esta función lee el contenido del archivo y lo descodifica como texto UTF-8
    antes de devolverlo
    :devolver: La función `extract_text_from_txt` lee el contenido de un archivo y devuelve el texto decodificado
    decodificado como una cadena en codificación UTF-8.
    """
    return file.read().decode('utf-8')

def extract_text_from_docx(file):
    """
    La función `extract_text_from_docx` toma un archivo .docx como entrada y devuelve el contenido de texto de
    del documento como una única cadena con párrafos separados por nuevas líneas.
    
    param archivo: El parámetro `file` en la función `extract_text_from_docx` se espera que sea la
    ruta a un archivo .docx del que desea extraer el texto
    :return: La función `extract_text_from_docx` toma un fichero como entrada, extrae el contenido de texto del
    el DocxDocument representado por el archivo, y devuelve el texto extraído como una única cadena con
    párrafos separados por caracteres de nueva línea.
    """
    doc = DocxDocument(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_pdf(file):
    """
    La función `extract_text_from_pdf` toma un archivo PDF como entrada, extrae el contenido de texto de cada
    página y devuelve el texto concatenado como una cadena.
    
    param archivo: Se espera que el parámetro `file` de la función `extract_text_from_pdf` sea un objeto file
    que representa un archivo PDF. La función lee el contenido del archivo PDF y extrae texto
    utilizando la biblioteca PyMuPDF (fitz). La función devuelve el texto extraído como una cadena
    :devolver: La función `extract_text_from_pdf` devuelve el contenido del texto extraído del archivo PDF.
    Si el texto extraído no está vacío (después de eliminar los espacios en blanco), devuelve el texto. En caso contrario
    devuelve `Ninguno`.
    """
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    text = "\n".join([page.get_text() for page in pdf_document])
    return text if text.strip() else None

def extract_text_from_image(file):
    """
    La función `extract_text_from_image` toma un archivo de imagen como entrada, utiliza pytesseract para extraer
    texto de la imagen, y devuelve el texto extraído como una cadena.
    
    param archivo: El parámetro `file` en la función `extract_text_from_image` se espera que sea un archivo
    que apunta a un archivo de imagen del que desea extraer el texto.
    :return: La función `extract_text_from_image` devuelve el texto extraído de la imagen utilizando
    pytesseract.
    """
    image = Image.open(file)
    return pytesseract.image_to_string(image)

def extract_text_from_excel(file):
    """
    La función `extract_text_from_excel` lee datos de un fichero Excel y devuelve el contenido de texto
    de cada hoja en una cadena formateada.
    
    param fichero: El parámetro `file` de la función `extract_text_from_excel` es la ruta al fichero Excel
    del que desea extraer el contenido de texto. Esta función lee el fichero Excel utilizando
    pd.read_excel` de la librería pandas, extrae el contenido de texto de cada hoja del fichero Excel,
    y devuelve
    :devuelve: La función `extract_text_from_excel` devuelve una cadena que contiene el contenido de texto
    extraído del fichero Excel. El contenido de texto incluye los datos de cada hoja del archivo Excel
    formateados como una cadena con los nombres de las hojas y sus respectivos datos. Si se produce un error durante el proceso de
    proceso de extracción, la función imprimirá un mensaje de error y devolverá `None`.
    """
    try:
        excel_data = pd.read_excel(file, sheet_name=None)
        text_content = []
        for sheet_name, df in excel_data.items():
            text_content.append(f"Hoja: {sheet_name}\n{df.to_string(index=False)}")
        return "\n\n".join(text_content)
    except Exception as e:
        print(f"Error al extraer texto del Excel: {e}")
        return None

def attach_thread():
    """Adjunta el hilo actual a la JVM."""
    lucene.getVMEnv().attachCurrentThread()

def get_index_writer():
    """
    La función `get_index_writer` crea y devuelve un objeto IndexWriter para escribir en un directorio de índices utilizando StandardAnalyzer.
    utilizando un StandardAnalyzer.
    :devolver: La función `get_index_writer()` devuelve un objeto `IndexWriter` configurado con un
    y configurado para escribir en el directorio de índices especificado por `INDEX_DIR`.
    """
    attach_thread()  # Adjuntamos el hilo actual
    analyzer = StandardAnalyzer()
    config = IndexWriterConfig(analyzer)
    index_path = File(INDEX_DIR).toPath()
    directory = FSDirectory.open(index_path)
    writer = IndexWriter(directory, config)
    return writer

def add_document(content, filename=None):
    """
    La función `add_document` añade un documento con contenido y un nombre de fichero opcional a un índice.
    
    param contenido: La función `add_document` se utiliza para añadir un documento a un índice. El parámetro `content
    representa el contenido textual del documento que se desea añadir al índice. Este contenido de
    contenido se almacenará en un campo llamado «content» en el documento Lucene.
    :param nombre_archivo: El parámetro `nombre_archivo` de la función `add_document` es un parámetro opcional
    que representa el nombre del fichero asociado al contenido que se añade al índice. Si se proporciona un
    nombre de fichero, se almacenará en el documento junto con el contenido.
    """
    writer = get_index_writer()
    doc = Document()
    doc.add(TextField("content", content, Field.Store.YES))
    if filename:
        doc.add(TextField("filename", filename, Field.Store.YES))
    writer.addDocument(doc)
    writer.commit()
    writer.close()

def search_documents(query_str):
    """
    La función `search_documents` busca documentos a partir de una cadena de consulta utilizando un índice y
    devuelve los resultados relevantes.
    
    param cadena_consulta: La función `buscar_documentos(cadena_consulta)` está diseñada para buscar documentos
    a partir de una cadena de consulta dada. Utiliza un índice creado con Apache Lucene para realizar la búsqueda.
    :devolver: La función `buscar_documentos` devuelve una lista de diccionarios, donde cada diccionario
    representa un documento que coincide con la consulta. Cada diccionario contiene el contenido del
    documento, el nombre del archivo y la puntuación que indica la relevancia del documento para la consulta.
    """
    attach_thread()  # Adjuntamos el hilo actual
    analyzer = StandardAnalyzer()
    index_path = File(INDEX_DIR).toPath()
    directory = FSDirectory.open(index_path)
    reader = DirectoryReader.open(directory)
    searcher = IndexSearcher(reader)
    query = QueryParser("content", analyzer).parse(query_str)
    hits = searcher.search(query, 10)
    results = []
    for hit in hits.scoreDocs:
        doc = searcher.storedFields().document(hit.doc)
        results.append({
            "content": doc.get("content"),
            "filename": doc.get("filename"),
            "score": hit.score
        })
    reader.close()
    return results

def index_postgres():
    """
    La función `index_postgres` se conecta a una base de datos PostgreSQL, recupera una lista de tablas en el esquema
    esquema «public», obtiene todas las filas de cada tabla y procesa los datos iterando a través de las filas y columnas.
    filas y columnas.
    """
    # Conecta a la base de datos usando host.docker.internal si la aplicación está en Docker
    conn = psycopg2.connect(
        host="host.docker.internal",
        port=5432,
        dbname="dbpostgrado",
        user="postgres",
        password="post123post"
    )
    cur = conn.cursor()
    
    # Obtén la lista de todas las tablas del esquema "public"
    cur.execute("""
        SELECT table_name 
        FROM information_schema.tables 
        WHERE table_schema = 'public' AND table_type = 'BASE TABLE'
    """)
    tables = cur.fetchall()

    for table in tables:
        table_name = table[0]
        print(f"Procesando la tabla: {table_name}")
        # Selecciona todas las filas de la tabla actual
        try:
            cur.execute(f"SELECT * FROM {table_name}")
            rows = cur.fetchall()
            # Obtén la descripción de las columnas para identificar los nombres, si es necesario
            col_names = [desc[0] for desc in cur.description]
            
            for row in rows:
                # Itera por cada columna de la fila
                for idx, cell in enumerate(row):
                    # Si deseas incluir el nombre de la columna junto con el contenido
                    col_name = col_names[idx]
                    # Asegúrate de convertir el contenido a cadena si no es None
                    if cell is not None:
                        content = f"{table_name}.{col_name}: {str(cell)}"
                        add_document(content)
        except Exception as e:
            print(f"Error al procesar la tabla {table_name}: {e}")
    
    cur.close()
    conn.close()

def index_file(file_path):
    """
    La función `index_file` lee el contenido de un fichero de texto y lo indexa.
    
    param `file_path: El parámetro `file_path` de la función `index_archivo` es una cadena que representa
    la ruta del fichero de texto que desea indexar. Esta función lee el contenido del fichero de texto
    ubicado en la `file_path` especificada y luego indexa ese contenido
    """

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            add_document(content)
        print(f"Archivo indexado: {file_path}")
    except Exception as e:
        print(f"Error al indexar el archivo {file_path}: {e}")
        raise

def index_folder(folder_path):
    """
    La función `index_folder` recorre recursivamente una carpeta, identifica los archivos de texto e indexa
    cada archivo.
    
    param folder_path: El parámetro `folder_path` de la función `index_folder` es una cadena que
    representa la ruta de la carpeta que desea indexar. Esta función utiliza el método `os.walk` para
    para recorrer todos los directorios y archivos dentro de la ruta especificada. Para cada archivo
    encontrado, comprueba si
    """
    
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            # Ajusta la validación de extensiones según necesites
            if file.lower().endswith(".txt"):
                file_path = os.path.join(root, file)
                index_file(file_path)

# Ruta de la interfaz principal
@app.route("/")
def home():
    return render_template("index.html")

# API para indexar la base de datos PostgreSQL
@app.route("/index_db", methods=["POST"])
def index_db():
    try:
        index_postgres()  # Función que indexa la BD
        flash("Base de datos indexada correctamente", "success")
    except Exception as e:
        flash(f"Error al indexar la base de datos: {str(e)}", "danger")
    return redirect(url_for("home"))

# API para indexar un documento individual
@app.route("/index", methods=["POST"])
def index_document():
    data = request.get_json()
    if not data or "content" not in data:
        return jsonify({
            "status": "error",
            "message": "No se proporcionó el contenido"
        }), 400
    try:
        add_document(data["content"])
        response = {
            "status": "success",
            "message": "Documento indexado correctamente",
            "document": data["content"]
        }
        return jsonify(response), 201
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500

# API para indexar múltiples documentos desde un formulario
@app.route("/index_path", methods=["POST"])
def index_path():
    if 'files' in request.files:
        files = request.files.getlist("files")
        for file in files:
            if file and allowed_file(file.filename):
                try:
                    filename = file.filename.lower()
                    text_content = None

                    if filename.endswith('.txt'):
                        text_content = extract_text_from_txt(file)
                    elif filename.endswith('.docx'):
                        text_content = extract_text_from_docx(file)
                    elif filename.endswith('.pdf'):
                        text_content = extract_text_from_pdf(file)
                    elif filename.endswith(('.png', '.jpg', '.jpeg')):
                        text_content = extract_text_from_image(file)
                    elif filename.endswith(('.xls', '.xlsx')):
                        text_content = extract_text_from_excel(file)

                    if text_content:
                        add_document(text_content, filename=file.filename)
                        print(f"Archivo indexado: {file.filename}")
                    else:
                        flash(f"El archivo {file.filename} no contiene texto legible.", "warning")

                except Exception as e:
                    print(f"Error al indexar {file.filename}: {e}")
                    flash(f"Error al indexar {file.filename}: {str(e)}", "danger")
        
        flash("Archivos indexados correctamente", "success")
    else:
        flash("No se proporcionaron archivos", "danger")
    
    return redirect(url_for("home"))


# Ruta de búsqueda: ahora renderiza la respuesta en HTML siempre.
@app.route("/search", methods=["GET"])
def search():
    query = request.args.get("q", "")
    results = []
    if query:
        results = search_documents(query)
    return render_template("index.html", results=results, query=query)

if __name__ == "__main__":
    if not os.path.exists(INDEX_DIR):
        os.makedirs(INDEX_DIR)
    app.run(host="0.0.0.0", port=5000)